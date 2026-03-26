"""
Resume Parser Service
----------------------
Extracts plain text from PDF and DOCX resume files.

Supports:
- PDF: via pdfplumber (handles multi-column layouts better than PyPDF2)
- DOCX: via python-docx
- Fallback: basic text extraction if primary method fails
"""

import os
from pathlib import Path

from app.utils.logger import get_logger

logger = get_logger(__name__)


class ResumeParser:
    """
    Extracts and cleans text from resume files.
    
    The extracted text is fed to the Claude evaluator.
    Quality of extraction directly impacts evaluation quality.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}

    def extract_text(self, file_path: str) -> str:
        """
        Main entry point — detects file type and delegates to the right parser.
        
        Args:
            file_path: Absolute or relative path to the resume file
            
        Returns:
            Cleaned plain text content of the resume
            
        Raises:
            ValueError: If file type is not supported
            RuntimeError: If parsing fails
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {self.SUPPORTED_EXTENSIONS}")

        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        logger.info(f"Parsing {ext} resume: {path.name}")

        if ext == ".pdf":
            return self._parse_pdf(str(path))
        elif ext in (".docx", ".doc"):
            return self._parse_docx(str(path))

    def _parse_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF using pdfplumber.
        
        pdfplumber handles:
        - Multi-column layouts
        - Tables
        - Headers/footers
        Better than PyPDF2 for modern resumes.
        """
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    else:
                        logger.warning(f"Empty page {page_num} in PDF (may be image-based)")

            full_text = "\n".join(text_parts)
            return self._clean_text(full_text)

        except ImportError:
            logger.warning("pdfplumber not available, falling back to pypdf2")
            return self._parse_pdf_fallback(file_path)
        except Exception as e:
            raise RuntimeError(f"PDF parsing failed: {e}") from e

    def _parse_pdf_fallback(self, file_path: str) -> str:
        """Fallback PDF parser using PyPDF2."""
        try:
            import PyPDF2

            text_parts = []
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text_parts.append(page.extract_text() or "")

            return self._clean_text("\n".join(text_parts))
        except Exception as e:
            raise RuntimeError(f"PDF fallback parsing failed: {e}") from e

    def _parse_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX using python-docx.
        
        Extracts:
        - Paragraph text (main body)
        - Table cell text (skills tables, etc.)
        - Text boxes are NOT extracted (limitation of python-docx)
        """
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            # Extract paragraph text
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract table text (skills and experience tables)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            return self._clean_text("\n".join(text_parts))

        except Exception as e:
            raise RuntimeError(f"DOCX parsing failed: {e}") from e

    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text for better LLM processing.
        
        - Remove excessive whitespace
        - Remove non-printable characters
        - Normalize line endings
        """
        import re

        # Remove non-printable characters (except newlines and tabs)
        text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", "", text)

        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # Remove excessive blank lines (max 2 consecutive)
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Strip trailing whitespace from each line
        lines = [line.rstrip() for line in text.split("\n")]
        text = "\n".join(lines)

        return text.strip()
