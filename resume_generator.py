"""
Resume Generator Service
-------------------------
Generates improved resumes in two formats:
1. ATS Format — Clean, single-column, keyword-optimized for parsing
2. Modern Format — Visually polished for human readers

Output formats: DOCX (editable) and PDF (final submission)

Uses python-docx for DOCX generation and docx2pdf for PDF conversion.
Falls back to reportlab if docx2pdf is unavailable (e.g., on Linux without LibreOffice).
"""

import os
import uuid
from typing import Dict, Any, Tuple
from datetime import datetime

from app.utils.logger import get_logger

logger = get_logger(__name__)

# Output directory for generated files
OUTPUT_DIR = "output"


class ResumeGenerator:
    """
    Generates improved resume files based on evaluation results.
    
    The generator applies:
    - Missing keyword insertions
    - Bullet point rewrites (if approved by user)
    - Updated professional summary
    - ATS-friendly formatting
    """

    def __init__(self):
        os.makedirs(OUTPUT_DIR, exist_ok=True)

    async def generate(
        self,
        resume_text: str,
        evaluation: Dict[str, Any],
        job_description: str,
    ) -> Tuple[str, str]:
        """
        Generate DOCX and PDF versions of the improved resume.
        
        Returns:
            Tuple of (docx_relative_path, pdf_relative_path)
        """
        session_id = uuid.uuid4().hex[:8]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = f"resume_{session_id}_{timestamp}"

        docx_path = os.path.join(OUTPUT_DIR, f"{base_name}.docx")
        pdf_path = os.path.join(OUTPUT_DIR, f"{base_name}.pdf")

        # Generate DOCX first (source of truth)
        self._generate_ats_docx(
            docx_path=docx_path,
            resume_text=resume_text,
            evaluation=evaluation,
        )

        # Convert DOCX → PDF
        self._convert_to_pdf(docx_path=docx_path, pdf_path=pdf_path)

        logger.info(f"Generated resume: {docx_path} | {pdf_path}")
        return docx_path, pdf_path

    def _generate_ats_docx(
        self,
        docx_path: str,
        resume_text: str,
        evaluation: Dict[str, Any],
    ):
        """
        Create an ATS-optimized DOCX resume.
        
        ATS Format Rules:
        - Single column layout
        - Standard fonts (Calibri/Times New Roman) — NOT decorative
        - No tables, text boxes, or headers/footers (ATS can't parse these)
        - Clear section headings with consistent formatting
        - Bullet points using standard characters (•)
        """
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re

        doc = Document()

        # ─── Page Margins (wider margins = more ATS-friendly) ─────────────────
        section = doc.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

        # ─── Parse resume into structured sections ────────────────────────────
        parsed = self._parse_resume_sections(resume_text)
        rewritten_bullets = {
            item["original"].strip(): item["rewritten"]
            for item in evaluation.get("rewritten_bullets", [])
            if evaluation.get("bullets_applied", False)
        }
        rewritten_summary = evaluation.get("rewritten_summary", "")

        # ─── Header: Name and Contact ─────────────────────────────────────────
        name = parsed.get("name", "Your Name")
        contact = parsed.get("contact", "")

        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_run.font.name = "Calibri"

        if contact:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(contact)
            contact_run.font.size = Pt(10)
            contact_run.font.name = "Calibri"

        # ─── ATS Optimization Note (hidden watermark in properties) ───────────
        # This doesn't affect ATS parsing — just metadata
        doc.core_properties.subject = "ATS-Optimized Resume"
        doc.core_properties.keywords = ", ".join(
            evaluation.get("keyword_optimization", {}).get("recommended_additions", [])[:10]
        )

        # ─── Professional Summary ─────────────────────────────────────────────
        summary_text = rewritten_summary or parsed.get("summary", "")
        if summary_text:
            self._add_section_heading(doc, "PROFESSIONAL SUMMARY")
            p = doc.add_paragraph(summary_text)
            p.runs[0].font.size = Pt(10.5)
            p.runs[0].font.name = "Calibri"

        # ─── Skills Section (keyword-optimized) ───────────────────────────────
        matched = evaluation.get("matched_skills", [])
        recommended = evaluation.get("keyword_optimization", {}).get("recommended_additions", [])
        all_skills = list(dict.fromkeys(matched + recommended))  # deduplicate, preserve order

        if all_skills:
            self._add_section_heading(doc, "CORE SKILLS")
            skills_text = " • ".join(all_skills[:20])  # ATS-friendly bullet separator
            p = doc.add_paragraph(skills_text)
            p.runs[0].font.size = Pt(10.5)
            p.runs[0].font.name = "Calibri"

        # ─── Experience Section ───────────────────────────────────────────────
        experience_blocks = parsed.get("experience", [])
        if experience_blocks:
            self._add_section_heading(doc, "PROFESSIONAL EXPERIENCE")
            for block in experience_blocks:
                # Job title + company
                job_para = doc.add_paragraph()
                job_run = job_para.add_run(block.get("title_line", ""))
                job_run.bold = True
                job_run.font.size = Pt(11)
                job_run.font.name = "Calibri"

                # Bullet points (with rewrites applied if approved)
                for bullet in block.get("bullets", []):
                    display_bullet = rewritten_bullets.get(bullet.strip(), bullet)
                    b_para = doc.add_paragraph(style="List Bullet")
                    b_run = b_para.add_run(display_bullet)
                    b_run.font.size = Pt(10.5)
                    b_run.font.name = "Calibri"

        # ─── Education Section ────────────────────────────────────────────────
        education = parsed.get("education", [])
        if education:
            self._add_section_heading(doc, "EDUCATION")
            for edu in education:
                edu_para = doc.add_paragraph()
                edu_run = edu_para.add_run(edu)
                edu_run.font.size = Pt(10.5)
                edu_run.font.name = "Calibri"

        # ─── ATS Warnings Footer ──────────────────────────────────────────────
        warnings = evaluation.get("ats_warnings", [])
        if warnings:
            # Add as comment-style note at bottom
            self._add_section_heading(doc, "OPTIMIZATION NOTES")
            for w in warnings[:3]:
                wp = doc.add_paragraph(f"• {w}")
                wp.runs[0].font.size = Pt(9)
                wp.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.save(docx_path)
        logger.info(f"DOCX saved: {docx_path}")

    def _add_section_heading(self, doc, text: str):
        """Add a bold, underlined, all-caps section heading (ATS standard)."""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  # Professional dark blue

        # Add bottom border (horizontal rule) for visual separation
        p_pr = p._p.get_or_add_pPr()
        p_bdr = p_pr.get_or_add_pBdr()
        bottom = p_bdr.get_or_add_bottom()
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F4E79")

    def _parse_resume_sections(self, resume_text: str) -> Dict[str, Any]:
        """
        Heuristically parse resume text into structured sections.
        
        This is intentionally simple — for production, consider spaCy NER
        or a dedicated Claude prompt for structured extraction.
        """
        import re

        lines = resume_text.split("\n")
        result = {
            "name": "",
            "contact": "",
            "summary": "",
            "experience": [],
            "education": [],
        }

        # First non-empty line is usually the name
        for line in lines[:5]:
            if line.strip() and not "@" in line and not "+" in line:
                result["name"] = line.strip()
                break

        # Contact info: lines with email, phone, LinkedIn
        contact_parts = []
        for line in lines[:10]:
            if any(x in line.lower() for x in ["@", "linkedin", "github", "+", "phone"]):
                contact_parts.append(line.strip())
        result["contact"] = " | ".join(contact_parts[:3])

        # Simple section detection
        section = None
        exp_block = None
        bullet_pattern = re.compile(r"^[\•\-\*\·▸►]\s*(.+)")

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            line_lower = stripped.lower()

            # Detect section headers
            if any(kw in line_lower for kw in ["experience", "work history", "employment"]):
                section = "experience"
                if exp_block:
                    result["experience"].append(exp_block)
                exp_block = {"title_line": "", "bullets": []}

            elif any(kw in line_lower for kw in ["education", "academic", "university", "degree"]):
                section = "education"
                if exp_block:
                    result["experience"].append(exp_block)
                    exp_block = None

            elif any(kw in line_lower for kw in ["summary", "objective", "profile", "about"]):
                section = "summary"

            else:
                # Add content to current section
                if section == "summary" and not result["summary"]:
                    result["summary"] = stripped

                elif section == "experience" and exp_block is not None:
                    bullet_match = bullet_pattern.match(stripped)
                    if bullet_match:
                        exp_block["bullets"].append(bullet_match.group(1))
                    elif not exp_block["title_line"] or (stripped.isupper() or any(c.isdigit() for c in stripped)):
                        if exp_block["title_line"]:
                            result["experience"].append(exp_block)
                        exp_block = {"title_line": stripped, "bullets": []}

                elif section == "education":
                    result["education"].append(stripped)

        if exp_block and exp_block["bullets"]:
            result["experience"].append(exp_block)

        return result

    def _convert_to_pdf(self, docx_path: str, pdf_path: str):
        """
        Convert DOCX to PDF.
        
        Primary: docx2pdf (requires Microsoft Word on Windows/Mac, or LibreOffice on Linux)
        Fallback: reportlab with basic formatting
        """
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            logger.info(f"PDF converted via docx2pdf: {pdf_path}")
        except Exception as e:
            logger.warning(f"docx2pdf failed ({e}), falling back to reportlab")
            self._generate_pdf_reportlab(docx_path, pdf_path)

    def _generate_pdf_reportlab(self, docx_path: str, pdf_path: str):
        """
        Generate PDF using reportlab as a fallback.
        Reads the DOCX text and creates a clean PDF layout.
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
            from reportlab.lib.colors import HexColor
            from docx import Document

            # Re-read DOCX for content
            doc = Document(docx_path)
            pdf_doc = SimpleDocTemplate(
                pdf_path,
                pagesize=letter,
                rightMargin=inch,
                leftMargin=inch,
                topMargin=0.75 * inch,
                bottomMargin=0.75 * inch,
            )

            styles = getSampleStyleSheet()
            story = []

            heading_style = ParagraphStyle(
                "heading",
                parent=styles["Normal"],
                fontSize=11,
                textColor=HexColor("#1F4E79"),
                fontName="Helvetica-Bold",
                spaceAfter=4,
            )
            body_style = ParagraphStyle(
                "body",
                parent=styles["Normal"],
                fontSize=10,
                fontName="Helvetica",
                spaceAfter=3,
            )
            name_style = ParagraphStyle(
                "name",
                parent=styles["Normal"],
                fontSize=18,
                fontName="Helvetica-Bold",
                alignment=TA_CENTER,
                spaceAfter=6,
            )

            for para in doc.paragraphs:
                if not para.text.strip():
                    story.append(Spacer(1, 6))
                    continue

                is_bold = any(run.bold for run in para.runs if run.text.strip())
                if is_bold and len(para.text) < 50:
                    if para.runs and para.runs[0].font.size and para.runs[0].font.size.pt >= 16:
                        story.append(Paragraph(para.text, name_style))
                    else:
                        story.append(Paragraph(para.text, heading_style))
                else:
                    story.append(Paragraph(para.text, body_style))

            pdf_doc.build(story)
            logger.info(f"PDF generated via reportlab: {pdf_path}")

        except Exception as e:
            logger.error(f"PDF generation failed entirely: {e}")
            # Create a minimal placeholder PDF
            self._create_placeholder_pdf(pdf_path)

    def _create_placeholder_pdf(self, pdf_path: str):
        """Create a minimal PDF with an error message as last resort."""
        try:
            from reportlab.pdfgen import canvas

            c = canvas.Canvas(pdf_path)
            c.drawString(100, 750, "Resume generated - Please open the DOCX file for full version.")
            c.save()
        except Exception:
            # If even this fails, create an empty file so the path exists
            with open(pdf_path, "wb") as f:
                f.write(b"%PDF-1.4\n")
